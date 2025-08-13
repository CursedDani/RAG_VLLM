import docx
from docx.document import Document as Document_compose
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import re
import argparse
import os
import logging
import pathlib
import requests

# --- Configuración ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def iter_block_items(parent):
    """
    Generador que itera sobre los párrafos y tablas de un documento, en orden.
    """
    if isinstance(parent, Document_compose):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Tipo de contenedor no soportado")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def main(input_path, output_path):
    
    """
    Función principal que construye un nuevo documento limpio
    a partir del documento original.
    """

    if not os.path.exists(input_path):
        logging.error(f"El archivo de entrada no existe: {input_path}")
        return

    logging.info(f"Cargando documento original: {input_path}")
    source_doc = docx.Document(input_path)
    
    dest_doc = docx.Document()
    
    logging.info("Iniciando proceso de reconstrucción limpia...")
    

    # --- FIN DE LA CORRECCIÓN ---

    for block in iter_block_items(source_doc):
        
        if isinstance(block, Paragraph):

            
            if block.text.strip():
                dest_doc.add_paragraph(block.text)

        elif isinstance(block, Table):
            logging.info("Procesando una tabla...")
            text_representation = []
            for row in block.rows:
                row_text = " | ".join([cell.text.strip().replace("\n", " ") for cell in row.cells])
                text_representation.append(row_text)
            
            full_text = "\n".join(text_representation)
            
            dest_doc.add_paragraph(full_text)
            dest_doc.add_paragraph()

    logging.info(f"Guardando nuevo documento limpio en: {output_path}")
    dest_doc.save(output_path)
    logging.info("✅ Proceso de reconstrucción completado exitosamente.")

APP_DIR = os.path.dirname(__file__)
INPUT_DIR = os.path.join(APP_DIR, "..", "data", "docs")
OUTPUT_DIR = os.path.join(APP_DIR, "..", "data", "docx")

filenames = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".docx")]
print(filenames)
for file in filenames:
    input_path = f'{INPUT_DIR}/{file}'
    output_path = f'{OUTPUT_DIR}/cleaned_{file}'
    main(input_path, output_path)

